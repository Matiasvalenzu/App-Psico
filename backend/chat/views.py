import re
import requests
from io import BytesIO
from django.conf import settings
from django.db.models import Count
from django.http import FileResponse
from pgvector.django import CosineDistance
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from pacientes.models import Paciente
from sesiones.embeddings import cosine_similarity, generate_text_embedding
from sesiones.models import Sesion, TranscripcionSegmento
from .models import ChatConversacion, ChatMensaje, InformeIA
from .serializers import (
    ChatConversacionSerializer,
    ChatConversacionListSerializer,
    ChatMensajeSerializer,
    InformeIASerializer,
)


class ChatConversacionViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action == "list":
            return ChatConversacionListSerializer
        return ChatConversacionSerializer

    def get_queryset(self):
        qs = (
            ChatConversacion.objects.filter(psicologo=self.request.user)
            .select_related("paciente", "psicologo")
            .prefetch_related("mensajes")
            .annotate(mensajes_count=Count("mensajes"))
        )
        paciente_id = self.request.query_params.get("paciente")
        if paciente_id:
            qs = qs.filter(paciente_id=paciente_id)
        return qs

    def perform_create(self, serializer):
        titulo = (serializer.validated_data.get("titulo") or "").strip()
        serializer.save(
            psicologo=self.request.user,
            titulo=titulo or "Nueva conversación",
        )

    @action(detail=True, methods=["post"])
    def enviar_mensaje(self, request, pk=None):
        conversacion = self.get_object()
        contenido = request.data.get("contenido", "")

        if not contenido:
            return Response(
                {"error": "El contenido es requerido"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        ChatMensaje.objects.create(
            conversacion=conversacion,
            rol=ChatMensaje.Rol.USER,
            contenido=contenido,
        )

        if conversacion.titulo in {"", "Nueva conversación"}:
            conversacion.titulo = self._titulo_desde_mensaje(contenido)
            conversacion.save(update_fields=["titulo", "updated_at"])
        else:
            conversacion.save(update_fields=["updated_at"])

        contexto_sesiones, fuentes = self._buscar_contexto(conversacion.paciente, contenido)
        contexto_paciente = self._contexto_paciente(conversacion.paciente)
        contexto = "\n\n".join(
            parte for parte in [contexto_paciente, contexto_sesiones] if parte.strip()
        )
        respuesta = self._consultar_deepseek(contenido, contexto)

        mensaje_asistente = ChatMensaje.objects.create(
            conversacion=conversacion,
            rol=ChatMensaje.Rol.ASSISTANT,
            contenido=respuesta,
            fuentes_json=fuentes,
        )

        return Response(ChatMensajeSerializer(mensaje_asistente).data)

    def _titulo_desde_mensaje(self, contenido):
        titulo = " ".join(contenido.split())[:80].strip()
        if not titulo:
            return "Nueva conversación"
        return f"{titulo}..." if len(contenido) > 80 else titulo

    def _buscar_contexto(self, paciente, consulta, top_k=15):
        query_embedding = generate_text_embedding(consulta)
        base_qs = TranscripcionSegmento.objects.filter(
            sesion__paciente=paciente
        ).select_related("sesion")

        try:
            segmentos = list(
                base_qs.exclude(embedding__isnull=True)
                .annotate(distance=CosineDistance("embedding", query_embedding))
                .order_by("distance")[: top_k * 5]
            )
        except Exception:
            segmentos = []

        if not segmentos:
            candidatos = list(base_qs.exclude(embedding__isnull=True)[:100])
            candidatos.sort(
                key=lambda seg: cosine_similarity(query_embedding, seg.embedding or []),
                reverse=True,
            )
            segmentos = candidatos[: top_k * 5]

        if not segmentos:
            segmentos = list(base_qs[: top_k * 5])

        seen_sessions = {}
        diversified = []
        for seg in segmentos:
            sid = seg.sesion_id
            if sid not in seen_sessions:
                seen_sessions[sid] = seg
                diversified.append(seg)
                if len(diversified) >= top_k:
                    break

        remaining = top_k - len(diversified)
        if remaining > 0:
            for seg in segmentos:
                if seg not in diversified:
                    diversified.append(seg)
                    if len(diversified) >= top_k:
                        break

        contexto = []
        fuentes = []
        for seg in diversified:
            fecha = seg.sesion.fecha_hora_inicio.strftime("%d/%m/%Y")
            if seg.sesion.origen == Sesion.Origen.DOCUMENTO_EXTERNO:
                etiqueta = f"Documento externo {fecha}"
                if seg.sesion.documento_nombre_original:
                    etiqueta = f"{etiqueta} - {seg.sesion.documento_nombre_original}"
                contexto.append(f"[{etiqueta}]: {seg.texto}")
            else:
                contexto.append(
                    f"[Sesión {fecha} - {seg.hablante}]: {seg.texto}"
                )
            fuentes.append(
                {
                    "segmento_id": seg.id,
                    "sesion_id": seg.sesion_id,
                    "fecha": fecha,
                    "origen": seg.sesion.origen,
                    "documento_nombre_original": seg.sesion.documento_nombre_original,
                    "hablante": seg.hablante,
                    "texto": seg.texto[:240],
                }
            )

        return "\n".join(contexto), fuentes

    def _contexto_paciente(self, paciente):
        campos = [
            ("Nombre", paciente.nombre_completo),
            ("Edad", f"{paciente.edad} años" if paciente.edad else ""),
            ("Sexo", paciente.get_sexo_display() if paciente.sexo else ""),
            ("Ocupación", paciente.ocupacion_laboral),
            ("Motivo de consulta", paciente.motivo_consulta),
            ("Objetivos de intervención", paciente.objetivos_intervencion),
            ("Diagnóstico sospechado", paciente.diagnostico_sospechado),
            ("Medicación actual", paciente.medicacion_actual),
            ("Frecuencia de atención", paciente.frecuencia_atencion),
            ("Derivación/interconsulta", paciente.derivacion_interconsulta),
        ]

        if paciente.riesgo_suicida:
            campos.append(("Alerta de riesgo", "Paciente marcado con riesgo suicida"))
        if paciente.ideacion_suicida_nivel:
            campos.append(
                ("Nivel de ideación suicida", str(paciente.ideacion_suicida_nivel))
            )

        lineas = [f"{label}: {valor}" for label, valor in campos if valor]
        if not lineas:
            return ""
        return "Contexto clínico registrado del paciente:\n" + "\n".join(lineas)

    def _consultar_deepseek(self, pregunta, contexto):
        if not settings.DEEPSEEK_API_KEY:
            return "API de DeepSeek no configurada. Configura DEEPSEEK_API_KEY en el entorno."

        system_prompt = (
            "Eres un asistente clínico de apoyo para psicólogos. "
            "Tu rol es ayudar al profesional a ordenar información, detectar patrones "
            "y preparar la siguiente sesión, sin reemplazar su criterio clínico. "
            "Responde siempre en español, con lenguaje prudente y profesional. "
            "No entregues diagnósticos definitivos; usa términos como hipótesis, "
            "elemento a explorar o sugerencia clínica cuando corresponda. "
            "Separa evidencia observada de inferencias. "
            "Si falta información, decláralo explícitamente. "
            "Cuando sea posible, cita fechas o fuentes de sesiones. "
            "Estructura la respuesta con estos bloques, omitiendo solo los que no apliquen: "
            "Resumen del proceso; Temas trabajados; Objetivos pendientes; "
            "Indicadores relevantes; Alertas o riesgos; Sugerencias para próxima sesión; "
            "Límites de la respuesta."
        )

        try:
            response = requests.post(
                f"{settings.DEEPSEEK_BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings.DEEPSEEK_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {
                            "role": "user",
                            "content": f"Contexto de sesiones:\n{contexto}\n\nPregunta: {pregunta}",
                        },
                    ],
                    "temperature": 0.3,
                    "max_tokens": 2000,
                },
                timeout=60,
            )
            data = response.json()
            respuesta = data["choices"][0]["message"]["content"]
            return respuesta
        except Exception as e:
            return f"Error al consultar DeepSeek: {str(e)}"


class InformeIAViewSet(viewsets.ModelViewSet):
    serializer_class = InformeIASerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = InformeIA.objects.filter(psicologo=self.request.user).select_related(
            "paciente", "psicologo", "sesion", "mensaje_origen"
        )
        paciente_id = self.request.query_params.get("paciente")
        if paciente_id:
            qs = qs.filter(paciente_id=paciente_id)
        return qs

    def perform_create(self, serializer):
        mensaje_origen = serializer.validated_data.get("mensaje_origen")
        tipo = serializer.validated_data.get("tipo") or InformeIA.Tipo.RESUMEN_CLINICO
        titulo = (serializer.validated_data.get("titulo") or "").strip()
        extra = {}

        if mensaje_origen:
            contenido = (serializer.validated_data.get("contenido") or "").strip()
            extra["contenido"] = contenido or mensaje_origen.contenido
            extra["fuentes_json"] = mensaje_origen.fuentes_json or []

        if not titulo:
            extra["titulo"] = self._titulo_por_tipo(tipo)

        serializer.save(psicologo=self.request.user, **extra)

    def _titulo_por_tipo(self, tipo):
        labels = dict(InformeIA.Tipo.choices)
        return labels.get(tipo, "Informe IA")

    def exportar_pdf(self, request, pk=None):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        informe = self.get_object()
        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 48

        def draw_line(text, font="Helvetica", size=10, leading=14):
            nonlocal y
            if y < 48:
                pdf.showPage()
                y = height - 48
            pdf.setFont(font, size)
            pdf.drawString(48, y, str(text)[:120])
            y -= leading

        draw_line(informe.titulo or self._titulo_por_tipo(informe.tipo), "Helvetica-Bold", 16, 22)
        draw_line(f"Paciente: {informe.paciente.nombre_completo}")
        draw_line(f"Tipo: {informe.get_tipo_display()}")
        draw_line(f"Fecha: {informe.created_at.strftime('%d/%m/%Y %H:%M')}")
        y -= 8
        draw_line("Contenido", "Helvetica-Bold", 12, 18)

        for raw_line in informe.contenido.splitlines() or [""]:
            words = raw_line.split()
            if not words:
                draw_line("")
                continue
            line = ""
            for word in words:
                if len(line) + len(word) > 95:
                    draw_line(line)
                    line = word
                else:
                    line = f"{line} {word}".strip()
            if line:
                draw_line(line)

        pdf.save()
        buffer.seek(0)
        return FileResponse(
            buffer,
            as_attachment=True,
            filename=f"{self._filename(informe)}.pdf",
        )

    def exportar_docx(self, request, pk=None):
        from docx import Document

        informe = self.get_object()
        document = Document()
        document.add_heading(informe.titulo or self._titulo_por_tipo(informe.tipo), level=1)
        document.add_paragraph(f"Paciente: {informe.paciente.nombre_completo}")
        document.add_paragraph(f"Tipo: {informe.get_tipo_display()}")
        document.add_paragraph(f"Fecha: {informe.created_at.strftime('%d/%m/%Y %H:%M')}")
        document.add_heading("Contenido", level=2)
        for raw_line in informe.contenido.splitlines() or [""]:
            document.add_paragraph(raw_line)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return FileResponse(
            buffer,
            as_attachment=True,
            filename=f"{self._filename(informe)}.docx",
        )

    def _filename(self, informe):
        base = informe.titulo or self._titulo_por_tipo(informe.tipo)
        clean = re.sub(r"[^A-Za-z0-9_-]+", "-", base.strip().lower()).strip("-")
        return clean or f"informe-ia-{informe.id}"
