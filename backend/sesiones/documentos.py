import os
import re
from io import BytesIO


ALLOWED_DOCUMENT_EXTENSIONS = {".txt", ".docx", ".pdf"}


class UnsupportedDocumentType(ValueError):
    pass


class DocumentTextExtractionError(ValueError):
    pass


def extract_text_from_uploaded_document(uploaded_file):
    filename = getattr(uploaded_file, "name", "")
    extension = os.path.splitext(filename)[1].lower()
    if extension not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise UnsupportedDocumentType(
            "Formato no soportado. Sube un archivo .txt, .docx o .pdf."
        )

    content = uploaded_file.read()
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)

    try:
        if extension == ".txt":
            text = _extract_txt(content)
        elif extension == ".docx":
            text = _extract_docx(content)
        else:
            text = _extract_pdf(content)
    except DocumentTextExtractionError:
        raise
    except Exception as exc:
        raise DocumentTextExtractionError(
            "No se pudo extraer texto del documento."
        ) from exc

    return _normalize_text(text)


def split_text_into_segments(text, max_chars=1600):
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    segments = []
    current = ""

    for block in blocks:
        if len(block) > max_chars:
            if current:
                segments.append(current)
                current = ""
            segments.extend(_split_long_block(block, max_chars))
            continue

        candidate = f"{current}\n\n{block}".strip() if current else block
        if len(candidate) <= max_chars:
            current = candidate
        else:
            segments.append(current)
            current = block

    if current:
        segments.append(current)

    return segments


def _extract_txt(content):
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentTextExtractionError("No se pudo leer el archivo TXT.")


def _extract_docx(content):
    from docx import Document

    document = Document(BytesIO(content))
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def _extract_pdf(content):
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    if reader.is_encrypted:
        raise DocumentTextExtractionError(
            "No se puede procesar un PDF protegido con contraseña."
        )

    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages)


def _split_long_block(block, max_chars):
    words = block.split()
    chunks = []
    current = ""

    for word in words:
        candidate = f"{current} {word}".strip() if current else word
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = word

    if current:
        chunks.append(current)

    return chunks


def _normalize_text(text):
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    normalized = []
    previous_blank = False

    for line in lines:
        if not line:
            if not previous_blank:
                normalized.append("")
            previous_blank = True
            continue
        normalized.append(line)
        previous_blank = False

    return "\n".join(normalized).strip()
